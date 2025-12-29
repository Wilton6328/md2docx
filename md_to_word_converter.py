"""
Markdown to Word Converter
高穩定性 Markdown 轉 Word 桌面應用程式
使用 CustomTkinter GUI + python-docx
確保使用 Word 內建樣式，無自定義樣式
"""

import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class MarkdownToWordConverter:
    """Markdown 轉 Word 核心轉換引擎"""
    
    def __init__(self):
        self.doc = None
        self.in_list = False
        self.list_level = 0
        
    def convert(self, md_path: str, docx_path: str) -> bool:
        """
        執行轉換
        
        Args:
            md_path: Markdown 檔案路徑
            docx_path: 輸出的 Word 檔案路徑
            
        Returns:
            bool: 轉換是否成功
        """
        try:
            # 讀取 Markdown 檔案（支援 UTF-8 編碼）
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 建立 Word 文件
            self.doc = Document()
            
            # 逐行解析並轉換
            lines = md_content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i]
                i = self._process_line(line, lines, i)
            
            # 儲存文件
            self.doc.save(docx_path)
            return True
            
        except FileNotFoundError:
            raise Exception(f"找不到檔案: {md_path}")
        except PermissionError:
            raise Exception(f"沒有權限寫入檔案: {docx_path}")
        except Exception as e:
            raise Exception(f"轉換過程發生錯誤: {str(e)}")
    
    def _process_line(self, line: str, lines: list, current_index: int) -> int:
        """
        處理單行 Markdown
        
        Args:
            line: 當前行內容
            lines: 所有行的列表
            current_index: 當前行索引
            
        Returns:
            int: 下一行的索引
        """
        # 處理空行
        if not line.strip():
            if not self.in_list:
                self.doc.add_paragraph()
            return current_index + 1
        
        # 處理標題 (H1-H6)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # 使用 Word 內建標題樣式
            para = self.doc.add_heading(level=level)
            para.text = text
            self.in_list = False
            return current_index + 1
        
        # 處理無序列表
        unordered_list_match = re.match(r'^(\s*)[*\-+]\s+(.+)$', line)
        if unordered_list_match:
            indent = len(unordered_list_match.group(1))
            text = unordered_list_match.group(2)
            self._add_list_item(text, indent, is_ordered=False)
            self.in_list = True
            return current_index + 1
        
        # 處理有序列表
        ordered_list_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_list_match:
            indent = len(ordered_list_match.group(1))
            text = ordered_list_match.group(2)
            self._add_list_item(text, indent, is_ordered=True)
            self.in_list = True
            return current_index + 1
        
        # 處理表格
        if '|' in line and line.strip().startswith('|'):
            return self._process_table(lines, current_index)
        
        # 處理程式碼區塊
        if line.strip().startswith('```'):
            return self._process_code_block(lines, current_index)
        
        # 處理引用
        quote_match = re.match(r'^>\s+(.+)$', line)
        if quote_match:
            text = quote_match.group(1)
            para = self.doc.add_paragraph(style='Quote')
            self._apply_inline_formatting(para, text)
            self.in_list = False
            return current_index + 1
        
        # 處理水平線
        if re.match(r'^[\-*_]{3,}$', line.strip()):
            # Word 沒有內建的水平線樣式，使用段落邊框
            para = self.doc.add_paragraph()
            self._add_horizontal_line(para)
            self.in_list = False
            return current_index + 1
        
        # 一般段落
        para = self.doc.add_paragraph(style='Normal')
        self._apply_inline_formatting(para, line)
        self.in_list = False
        return current_index + 1
    
    def _add_list_item(self, text: str, indent: int, is_ordered: bool):
        """
        新增列表項目
        
        Args:
            text: 列表項目文字
            indent: 縮排層級
            is_ordered: 是否為有序列表
        """
        # 使用 Word 內建列表樣式
        style = 'List Number' if is_ordered else 'List Bullet'
        para = self.doc.add_paragraph(style=style)
        
        # 設定縮排層級
        level = indent // 2  # 每 2 個空格為一層
        para.paragraph_format.left_indent = Inches(0.5 * level)
        
        # 套用行內格式
        self._apply_inline_formatting(para, text)
    
    def _apply_inline_formatting(self, paragraph, text: str):
        """
        套用行內格式（粗體、斜體、程式碼等）
        
        Args:
            paragraph: Word 段落物件
            text: 要處理的文字
        """
        # 處理行內程式碼、粗體、斜體
        parts = self._parse_inline_formatting(text)
        
        for part_text, is_bold, is_italic, is_code in parts:
            run = paragraph.add_run(part_text)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True
            if is_code:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
    
    def _parse_inline_formatting(self, text: str) -> list:
        """
        解析行內格式
        
        Returns:
            list: [(文字, 是否粗體, 是否斜體, 是否程式碼), ...]
        """
        result = []
        current_pos = 0
        
        # 使用正則表達式找出所有格式標記
        # 優先處理程式碼，避免與其他格式衝突
        pattern = r'(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)'
        
        for match in re.finditer(pattern, text):
            # 添加前面的普通文字
            if match.start() > current_pos:
                plain_text = text[current_pos:match.start()]
                result.append((plain_text, False, False, False))
            
            # 處理格式化文字
            matched_text = match.group(0)
            
            # 行內程式碼
            if matched_text.startswith('`') and matched_text.endswith('`'):
                code_text = matched_text[1:-1]
                result.append((code_text, False, False, True))
            
            # 粗體 + 斜體
            elif matched_text.startswith('***') and matched_text.endswith('***'):
                content = matched_text[3:-3]
                result.append((content, True, True, False))
            
            # 粗體
            elif (matched_text.startswith('**') and matched_text.endswith('**')) or \
                 (matched_text.startswith('__') and matched_text.endswith('__')):
                content = matched_text[2:-2]
                result.append((content, True, False, False))
            
            # 斜體
            elif (matched_text.startswith('*') and matched_text.endswith('*')) or \
                 (matched_text.startswith('_') and matched_text.endswith('_')):
                content = matched_text[1:-1]
                result.append((content, False, True, False))
            
            current_pos = match.end()
        
        # 添加剩餘的普通文字
        if current_pos < len(text):
            result.append((text[current_pos:], False, False, False))
        
        return result if result else [(text, False, False, False)]
    
    def _process_table(self, lines: list, start_index: int) -> int:
        """
        處理 Markdown 表格
        
        Returns:
            int: 下一行的索引
        """
        table_lines = []
        i = start_index
        
        # 收集所有表格行
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i])
            i += 1
        
        if len(table_lines) < 2:
            return i
        
        # 解析表格
        rows = []
        max_cols = 0
        for line in table_lines:
            # 跳過分隔線
            if re.match(r'^\|[\s\-:]+\|$', line.strip()):
                continue
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:  # 只添加非空行
                rows.append(cells)
                max_cols = max(max_cols, len(cells))
        
        if not rows or max_cols == 0:
            return i
        
        # 標準化所有行到相同列數
        for row in rows:
            while len(row) < max_cols:
                row.append('')  # 補齊空儲存格
        
        # 建立 Word 表格
        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'  # 使用 Word 內建表格樣式
        
        # 填入資料
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):  # 安全檢查
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_data
                    
                    # 第一列設為粗體（標題列）
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        self.in_list = False
        return i
    
    def _process_code_block(self, lines: list, start_index: int) -> int:
        """
        處理程式碼區塊
        
        Returns:
            int: 下一行的索引
        """
        i = start_index + 1
        code_lines = []
        
        # 收集程式碼內容
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # 建立程式碼段落
        code_text = '\n'.join(code_lines)
        para = self.doc.add_paragraph(style='Normal')
        run = para.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        # 添加背景色（淺灰色）
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F0F0')
        para._element.get_or_add_pPr().append(shading_elm)
        
        self.in_list = False
        return i + 1
    
    def _add_horizontal_line(self, paragraph):
        """新增水平線"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)


class App(ctk.CTk):
    """主應用程式視窗"""
    
    def __init__(self):
        super().__init__()
        
        # 視窗設定
        self.title("Markdown to Word Converter")
        self.geometry("700x400")
        self.resizable(False, False)
        
        # 設定主題
        ctk.set_appearance_mode("system")
        ctk.set_default_color_theme("blue")
        
        # 變數
        self.file_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.status_text = tk.StringVar(value="就緒")
        
        # 建立 UI
        self._create_widgets()
        
        # 轉換器實例
        self.converter = MarkdownToWordConverter()
    
    def _create_widgets(self):
        """建立所有 UI 元件"""
        
        # 標題
        title_label = ctk.CTkLabel(
            self,
            text="📝 Markdown to Word Converter",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=20)
        
        # 檔案選擇框架
        file_frame = ctk.CTkFrame(self)
        file_frame.pack(pady=10, padx=40, fill="x")
        
        # 輸入檔案
        input_label = ctk.CTkLabel(file_frame, text="Markdown 檔案:", font=ctk.CTkFont(size=14))
        input_label.grid(row=0, column=0, padx=10, pady=10, sticky="w")
        
        self.file_entry = ctk.CTkEntry(
            file_frame,
            textvariable=self.file_path,
            width=400,
            placeholder_text="請選擇 .md 檔案"
        )
        self.file_entry.grid(row=0, column=1, padx=10, pady=10)
        
        browse_btn = ctk.CTkButton(
            file_frame,
            text="瀏覽",
            command=self.browse_file,
            width=100
        )
        browse_btn.grid(row=0, column=2, padx=10, pady=10)
        
        # 輸出檔案
        output_label = ctk.CTkLabel(file_frame, text="輸出位置:", font=ctk.CTkFont(size=14))
        output_label.grid(row=1, column=0, padx=10, pady=10, sticky="w")
        
        self.output_entry = ctk.CTkEntry(
            file_frame,
            textvariable=self.output_path,
            width=400,
            placeholder_text="自動生成在相同目錄"
        )
        self.output_entry.grid(row=1, column=1, padx=10, pady=10)
        
        output_browse_btn = ctk.CTkButton(
            file_frame,
            text="瀏覽",
            command=self.browse_output,
            width=100
        )
        output_browse_btn.grid(row=1, column=2, padx=10, pady=10)
        
        # 轉換按鈕
        self.convert_btn = ctk.CTkButton(
            self,
            text="🔄 開始轉換",
            command=self.convert_file,
            font=ctk.CTkFont(size=16, weight="bold"),
            height=50,
            width=200
        )
        self.convert_btn.pack(pady=30)
        
        # 狀態顯示
        status_frame = ctk.CTkFrame(self)
        status_frame.pack(pady=10, padx=40, fill="x")
        
        status_label = ctk.CTkLabel(
            status_frame,
            text="狀態:",
            font=ctk.CTkFont(size=14)
        )
        status_label.pack(side="left", padx=10)
        
        self.status_display = ctk.CTkLabel(
            status_frame,
            textvariable=self.status_text,
            font=ctk.CTkFont(size=14),
            text_color="green"
        )
        self.status_display.pack(side="left", padx=10)
    
    def browse_file(self):
        """選擇 Markdown 檔案"""
        filename = filedialog.askopenfilename(
            title="選擇 Markdown 檔案",
            filetypes=[("Markdown 檔案", "*.md"), ("所有檔案", "*.*")]
        )
        
        if filename:
            self.file_path.set(filename)
            
            # 自動設定輸出路徑
            if not self.output_path.get():
                base_path = Path(filename)
                output_filename = base_path.stem + ".docx"
                output_full_path = base_path.parent / output_filename
                self.output_path.set(str(output_full_path))
    
    def browse_output(self):
        """選擇輸出位置"""
        filename = filedialog.asksaveasfilename(
            title="選擇輸出位置",
            defaultextension=".docx",
            filetypes=[("Word 文件", "*.docx"), ("所有檔案", "*.*")]
        )
        
        if filename:
            self.output_path.set(filename)
    
    def convert_file(self):
        """執行轉換"""
        # 驗證輸入
        input_file = self.file_path.get()
        if not input_file:
            messagebox.showerror("錯誤", "請選擇要轉換的 Markdown 檔案！")
            return
        
        if not os.path.exists(input_file):
            messagebox.showerror("錯誤", "選擇的檔案不存在！")
            return
        
        # 設定輸出路徑
        output_file = self.output_path.get()
        if not output_file:
            base_path = Path(input_file)
            output_filename = base_path.stem + ".docx"
            output_file = str(base_path.parent / output_filename)
            self.output_path.set(output_file)
        
        # 更新狀態
        self.status_text.set("轉換中...")
        self.status_display.configure(text_color="orange")
        self.convert_btn.configure(state="disabled")
        self.update()
        
        try:
            # 執行轉換
            self.converter.convert(input_file, output_file)
            
            # 成功
            self.status_text.set("✓ 轉換完成！")
            self.status_display.configure(text_color="green")
            messagebox.showinfo("成功", f"轉換完成！\n檔案已儲存至：\n{output_file}")
            
            # 詢問是否開啟檔案
            if messagebox.askyesno("開啟檔案", "是否要開啟轉換後的 Word 檔案？"):
                os.startfile(output_file)
        
        except Exception as e:
            self.status_text.set("✗ 轉換失敗")
            self.status_display.configure(text_color="red")
            messagebox.showerror("錯誤", f"轉換失敗：\n{str(e)}")
        
        finally:
            self.convert_btn.configure(state="normal")


def main():
    """主程式進入點"""
    app = App()
    app.mainloop()


if __name__ == "__main__":
    main()
