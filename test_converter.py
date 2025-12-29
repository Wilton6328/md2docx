"""
測試腳本：驗證轉換器核心功能
"""

from md_to_word_converter import MarkdownToWordConverter
import os

def test_conversion():
    """測試基本轉換功能"""
    print("=" * 50)
    print("Markdown to Word Converter - 功能測試")
    print("=" * 50)
    
    # 初始化轉換器
    converter = MarkdownToWordConverter()
    
    # 設定檔案路徑
    input_file = "test_sample.md"
    output_file = "test_output.docx"
    
    # 檢查輸入檔案
    if not os.path.exists(input_file):
        print(f"[X] 錯誤：找不到測試檔案 {input_file}")
        return False
    
    print(f"\n輸入檔案: {input_file}")
    print(f"輸出檔案: {output_file}")
    print("\n開始轉換...")
    
    try:
        # 執行轉換
        converter.convert(input_file, output_file)
        
        # 驗證輸出
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file)
            print(f"\n[OK] 轉換成功！")
            print(f"檔案大小: {file_size:,} bytes")
            
            # 驗證文件可正常開啟
            from docx import Document
            doc = Document(output_file)
            print(f"段落數量: {len(doc.paragraphs)}")
            print(f"表格數量: {len(doc.tables)}")
            
            return True
        else:
            print(f"\n[X] 錯誤：輸出檔案未建立")
            return False
            
    except Exception as e:
        print(f"\n[X] 轉換失敗：{str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_conversion()
    print("\n" + "=" * 50)
    if success:
        print("[OK] 測試通過！")
    else:
        print("[FAILED] 測試失敗！")
    print("=" * 50)
